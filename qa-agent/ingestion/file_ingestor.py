"""
File-based PRD ingestor for Phase 2.

Accepts uploaded files (Excel, Word, PDF, Markdown, plain text) and converts
them to chunks ready for embedding.

Supported formats:
  .xlsx          — Excel (each sheet becomes a section)
  .docx          — Word document (headings become sections)
  .pdf           — PDF (page-based chunking)
  .md            — Markdown (section-based, same as GitLab/Confluence)
  .txt           — Plain text (no headings, single section)

Tabular content from every format is rendered through ingestion/markdown_table.py
so the chunker's table handling (row-level splitting, header carry-over) engages
uniformly. Images are reduced to a text marker naming them — the alt text and
filename are usually the only searchable trace a diagram leaves.
"""
import io
import logging
import re
import tempfile
import os

from ingestion.doc_classify import classify as _classify_doc
from ingestion.markdown_table import rows_to_markdown

logger = logging.getLogger(__name__)

# Guard against a spreadsheet or PDF large enough to exhaust memory during ingest.
MAX_SHEET_ROWS = int(os.environ.get("QA_INGEST_MAX_SHEET_ROWS", "20000"))
# Above this, .xlsx is read in streaming mode — see _excel_to_markdown.
MAX_WORKBOOK_BYTES = int(os.environ.get("QA_INGEST_MAX_WORKBOOK_MB", "25")) * 1_000_000
# OCR for scanned PDF pages. Off by default — needs the tesseract binary in the image.
PDF_OCR_ENABLED = os.environ.get("QA_INGEST_PDF_OCR", "0").strip() not in ("", "0", "false", "False")
OCR_DPI = int(os.environ.get("QA_INGEST_OCR_DPI", "200"))


# ─── Format converters ─────────────────────────────────────────────────────────

def _fill_merged_cells(ws) -> dict[tuple[int, int], str]:
    """
    Map every cell covered by a merged range to that range's top-left value.

    Merged cells are normal in spec sheets (one "Module" cell spanning several
    requirement rows). openpyxl reports the covered cells as None, so without
    this a row split into its own chunk loses the value entirely.
    """
    filled: dict[tuple[int, int], str] = {}
    merged = getattr(ws, "merged_cells", None)
    for rng in getattr(merged, "ranges", ()):
        try:
            anchor = ws.cell(row=rng.min_row, column=rng.min_col).value
        except Exception:
            continue
        if anchor is None:
            continue
        value = str(anchor).strip()
        for r in range(rng.min_row, rng.max_row + 1):
            for c in range(rng.min_col, rng.max_col + 1):
                filled[(r, c)] = value
    return filled


def _excel_to_markdown(content: bytes, filename: str) -> str:
    """Convert Excel file to markdown text (each sheet → H2 section)."""
    try:
        import openpyxl
    except ImportError:
        raise ImportError("openpyxl required for Excel ingestion: pip install openpyxl")

    # read_only=True streams rows and keeps memory flat, but does not expose
    # merged-cell ranges, so merged spec columns cannot be filled down. Normal PRD
    # workbooks are small enough to load fully; a large one falls back to streaming
    # rather than risking the worker's memory, and loses only the merge fill
    # (_fill_merged_cells no-ops when the attribute is absent).
    streaming = len(content) > MAX_WORKBOOK_BYTES
    if streaming:
        logger.warning(
            "Workbook %r is %.1f MB (> %.1f MB): reading in streaming mode. Merged cells "
            "will not be filled down; covered cells read as empty.",
            filename, len(content) / 1e6, MAX_WORKBOOK_BYTES / 1e6,
        )
    wb = openpyxl.load_workbook(
        io.BytesIO(content), read_only=streaming, data_only=True
    )
    parts = [f"# {filename}\n"]

    try:
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            merged = _fill_merged_cells(ws)

            rows: list[list[str]] = []
            for r, row in enumerate(ws.iter_rows(), start=1):
                if len(rows) >= MAX_SHEET_ROWS:
                    logger.warning(
                        "Sheet %r truncated at %s rows (QA_INGEST_MAX_SHEET_ROWS)",
                        sheet_name, MAX_SHEET_ROWS,
                    )
                    break
                cells = []
                for c, cell in enumerate(row, start=1):
                    value = merged.get((r, c))
                    if value is None:
                        value = "" if cell.value is None else str(cell.value).strip()
                    cells.append(value)
                # Skip completely empty rows — but never let a leading blank row
                # decide which row is the header (the old `i == 0` check did,
                # and a sheet starting with a blank row got no header at all).
                if any(cells):
                    rows.append(cells)

            if not rows:
                logger.warning(
                    "Sheet %r produced no rows. If it is formula-driven, the workbook may "
                    "have no cached values (data_only=True returns None until Excel saves it).",
                    sheet_name,
                )
                continue

            parts.append(f"\n## {sheet_name}\n")
            parts.append(rows_to_markdown(rows))
    finally:
        wb.close()

    return "\n".join(parts)


_DRAWING_NS = {
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a":  "http://schemas.openxmlformats.org/drawingml/2006/main",
}


def _docx_paragraph_markdown(para) -> str:
    """One Word paragraph → a markdown line, mapping heading styles to '#' levels."""
    text = para.text.strip()
    images = _docx_paragraph_images(para)
    if images:
        text = (text + " " if text else "") + " ".join(images)
    if not text:
        return ""

    style = ""
    try:
        style = para.style.name if para.style else ""
    except Exception:
        style = ""

    if "Heading 1" in style:
        return f"\n# {text}"
    if "Heading 2" in style:
        return f"\n## {text}"
    if "Heading 3" in style:
        return f"\n### {text}"
    if style.startswith("List"):
        return f"- {text}"
    return text


def _docx_paragraph_images(para) -> list[str]:
    """
    Text markers for inline images, from the drawing's name/description.

    Without this a diagram leaves no trace in the index at all.
    """
    markers: list[str] = []
    try:
        for props in para._p.iterfind(".//wp:docPr", _DRAWING_NS):
            label = (props.get("descr") or props.get("name") or "").strip()
            markers.append(f"[Image: {label}]" if label else "[Image]")
    except Exception:                     # unusual drawing markup must not stop ingest
        return markers
    return markers


def _docx_table_markdown(table) -> str:
    """Word table → markdown pipe table, resolving merged cells by repeated text."""
    rows: list[list[str]] = []
    for row in table.rows:
        cells: list[str] = []
        for cell in row.cells:
            # python-docx repeats the same cell object across a merged span, so the
            # value lands in every covered column — which is what we want here.
            cells.append(" ".join(p.text.strip() for p in cell.paragraphs if p.text.strip()))
        rows.append(cells)
    return rows_to_markdown(rows)


def _docx_to_markdown(content: bytes, filename: str) -> str:
    """
    Convert Word document to markdown (heading styles → # headings, tables → pipe tables).

    Walks body content in document order. `doc.paragraphs` alone returns only
    `w:p` children of `w:body` — paragraphs inside `w:tbl` are excluded — so a
    .docx whose requirements live in tables previously ingested as headings only.
    """
    try:
        from docx import Document
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError:
        raise ImportError("python-docx required for Word ingestion: pip install python-docx")

    doc = Document(io.BytesIO(content))
    lines: list[str] = [f"# {filename}\n"]

    if hasattr(doc, "iter_inner_content"):
        body = doc.iter_inner_content()
    else:                                  # python-docx < 1.1 — tables after paragraphs
        body = list(doc.paragraphs) + list(doc.tables)
        logger.warning(
            "python-docx has no iter_inner_content(); tables will be appended after the "
            "body text instead of appearing in document order."
        )

    n_tables = 0
    for item in body:
        if isinstance(item, Paragraph):
            line = _docx_paragraph_markdown(item)
            if line:
                lines.append(line)
        elif isinstance(item, Table):
            markdown = _docx_table_markdown(item)
            if markdown:
                n_tables += 1
                lines.append("")
                lines.append(markdown)
                lines.append("")

    logger.info("docx %r: %s tables extracted", filename, n_tables)
    return "\n".join(lines)


def _pdf_page_markdown(page) -> str:
    """
    One PDF page → prose followed by its tables as markdown.

    Tables are lifted out and the prose is re-extracted with the table regions
    masked, so table cells are not also emitted as a jumbled run of text.
    `extract_text()` alone flattens a table into whitespace-separated values with
    no column association — which is what this file did before.
    """
    tables: list[str] = []
    boxes: list[tuple[float, float, float, float]] = []

    try:
        for found in page.find_tables():
            rows = found.extract() or []
            markdown = rows_to_markdown([[c or "" for c in row] for row in rows])
            if markdown:
                tables.append(markdown)
                boxes.append(found.bbox)
    except Exception as exc:
        logger.warning("PDF table extraction failed on one page: %s", exc)

    text = ""
    try:
        if boxes:
            def _outside_tables(obj) -> bool:
                x = (obj.get("x0", 0) + obj.get("x1", 0)) / 2
                y = (obj.get("top", 0) + obj.get("bottom", 0)) / 2
                return not any(x0 <= x <= x1 and t <= y <= b for x0, t, x1, b in boxes)

            text = page.filter(_outside_tables).extract_text() or ""
        else:
            text = page.extract_text() or ""
    except Exception as exc:
        logger.warning("PDF text extraction fell back to unfiltered text: %s", exc)
        text = page.extract_text() or ""

    blocks = [text.strip()] if text.strip() else []
    blocks.extend(tables)
    return "\n\n".join(blocks)


def _ocr_page(page) -> str:
    """
    Last-resort text for a page pdfplumber found nothing in (a scanned image).

    Opt-in via QA_INGEST_PDF_OCR=1 and gated on the optional dependencies being
    importable, because OCR needs the tesseract binary in the image — neither is
    installed by default, and a missing one must degrade to today's behaviour
    rather than fail the ingest.
    """
    if not PDF_OCR_ENABLED:
        return ""
    try:
        import pytesseract                       # noqa: F401  (optional dependency)
        from PIL import Image                    # noqa: F401  (optional dependency)
    except ImportError:
        logger.warning(
            "QA_INGEST_PDF_OCR=1 but pytesseract/Pillow are not installed — "
            "scanned pages stay empty. Add both plus the tesseract binary to the image."
        )
        return ""
    try:
        image = page.to_image(resolution=OCR_DPI).original
        return (pytesseract.image_to_string(image) or "").strip()
    except Exception as exc:
        logger.warning("OCR failed on one page: %s", exc)
        return ""


def _pdf_to_text(content: bytes, filename: str) -> str:
    """Convert PDF to text (page-by-page, each page → H2 section)."""
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber required for PDF ingestion: pip install pdfplumber")

    parts = [f"# {filename}\n"]
    empty_pages = 0
    ocr_pages = 0
    total_pages = 0

    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            total_pages += 1
            body = _pdf_page_markdown(page).strip()
            if not body:
                body = _ocr_page(page)
                if body:
                    ocr_pages += 1
            if body:
                parts.append(f"\n## Page {i}\n\n{body}")
            else:
                empty_pages += 1

    if ocr_pages:
        logger.info("PDF %r: %s page(s) recovered by OCR", filename, ocr_pages)
    if empty_pages:
        logger.warning(
            "PDF %r: %s of %s pages yielded no extractable text and index as nothing. "
            "Scanned or image-only PDFs need OCR — %s",
            filename, empty_pages, total_pages,
            "OCR ran but found no text." if PDF_OCR_ENABLED
            else "set QA_INGEST_PDF_OCR=1 (needs pytesseract, Pillow and the tesseract binary).",
        )

    return "\n".join(parts)


def _text_to_markdown(content: bytes, filename: str) -> str:
    """Plain text / markdown — decode and prepend title heading if not present."""
    text = content.decode("utf-8", errors="replace").strip()
    if not text.startswith("#"):
        return f"# {filename}\n\n{text}"
    return text


# ─── Main entry ───────────────────────────────────────────────────────────────

def convert_file_to_markdown(
    filename: str,
    content: bytes,
) -> str:
    """
    Convert an uploaded file to markdown/plain text for chunking.

    Args:
        filename: original filename (used to detect format and as doc title)
        content:  raw file bytes

    Returns:
        Markdown string ready for chunk_document()

    Raises:
        ValueError if format is unsupported
        ImportError if required library is not installed
    """
    ext = os.path.splitext(filename.lower())[1]
    # Clean up filename for use as title
    title = re.sub(r'\.[^.]+$', '', filename)  # strip extension

    if ext == ".xls":
        raise ValueError(
            "Legacy .xls format is not supported. Please convert to .xlsx (Excel 2007+) and re-upload."
        )
    if ext == ".xlsx":
        return _excel_to_markdown(content, title)
    elif ext == ".docx":
        return _docx_to_markdown(content, title)
    elif ext == ".pdf":
        return _pdf_to_text(content, title)
    elif ext in (".md", ".txt", ".text"):
        return _text_to_markdown(content, title)
    else:
        raise ValueError(
            f"Unsupported file format: {ext!r}. "
            "Supported: .xlsx, .xls, .docx, .pdf, .md, .txt"
        )


def ingest_file(filename: str, content: bytes, source_label: str = "", embed_fn=None) -> list[dict]:
    """
    Full ingestion flow for a single uploaded file.

    Converts file to markdown, then chunks it properly to stay within
    the embedding model's token limits.

    Args:
        filename:     original filename (determines format)
        content:      raw file bytes
        source_label: optional human label; defaults to filename without extension
    """
    from ingestion.chunker import chunk_document

    label = source_label or re.sub(r'\.[^.]+$', '', filename)
    # Namespace uploads so user-controlled labels cannot collide with confluence:/gitlab: IDs
    safe = re.sub(r"[^\w\-. ]", "_", label.strip())[:200] or "document"
    source_id = f"file:upload:{safe}"
    ext       = os.path.splitext(filename.lower())[1].lstrip(".")

    full_text = convert_file_to_markdown(filename, content)

    # Chunk the document properly instead of returning as one oversized block
    raw_chunks = chunk_document(full_text, source_id, embed_fn=embed_fn)

    if not raw_chunks:
        logger.warning(f"File {filename!r} produced no chunks after processing")
        return []

    logger.info(f"File {filename!r} → {len(raw_chunks)} chunks (source_id={source_id!r})")
    return [
        {
            "source_id":       source_id,
            "source_type":     f"file_{ext}" if ext else "file",
            "source_version":  "uploaded",
            "doc_title":       label,
            "doc_url":         None,
            "section_heading": c.get("section_heading", label),
            "chunk_text":      c["chunk_text"],
            "parent_text":     c.get("parent_text"),
            "chunk_type":      c.get("chunk_type"),
            "doc_type":        _classify_doc(label),
            "chunk_index":     i,
        }
        for i, c in enumerate(raw_chunks)
    ]
