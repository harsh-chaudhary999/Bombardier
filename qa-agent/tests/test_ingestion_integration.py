"""
Integration checks against the REAL ingestion libraries.

Every other ingestion test runs on stubs or duck-typed stand-ins, so they verify
our logic against our *model* of html2text, python-docx, openpyxl and pdfplumber.
That model is where the bugs hide: `doc.paragraphs` silently excluding table
paragraphs, and html2text discarding CDATA, both looked correct on a stub.

These build real files in memory and assert on real output. Each class skips
when its library is absent, so the suite still runs on a bare host — but inside
the container, where requirements.txt is installed, they all execute.

Runs under the documented host-side command:
    cd qa-agent && PYTHONPATH=. python3 -m unittest discover -s tests -p 'test_*.py'
"""
import importlib.util
import io
import unittest


def _has(module: str) -> bool:
    try:
        return importlib.util.find_spec(module) is not None
    except (ImportError, ValueError):
        return False


@unittest.skipUnless(_has("html2text"), "html2text not installed")
class ConfluenceRealHtml2TextTests(unittest.TestCase):
    """The CDATA and image losses were invisible to a stubbed renderer."""

    def setUp(self):
        from ingestion.confluence_html import storage_to_markdown
        self.convert = storage_to_markdown

    def test_code_macro_body_survives_real_html2text(self):
        html = (
            "<p>Request:</p>"
            '<ac:structured-macro ac:name="code">'
            '<ac:parameter ac:name="language">json</ac:parameter>'
            '<ac:plain-text-body><![CDATA[{"amount": 0, "currency": "XXX"}]]>'
            "</ac:plain-text-body></ac:structured-macro>"
        )
        out = self.convert(html)
        self.assertIn('{"amount": 0, "currency": "XXX"}', out)
        self.assertIn("```json", out)

    def test_placeholders_are_never_mangled_by_the_renderer(self):
        """The whole design rests on html2text passing our tokens through intact."""
        html = (
            "<p>Before</p>"
            '<ac:structured-macro ac:name="code">'
            "<ac:plain-text-body><![CDATA[TOKEN_BODY_1]]></ac:plain-text-body>"
            "</ac:structured-macro><p>After</p>"
        )
        out = self.convert(html)
        self.assertIn("TOKEN_BODY_1", out)
        self.assertNotIn("ZQX", out, "no placeholder may survive into the output")

    def test_real_table_becomes_pipe_rows(self):
        html = (
            "<table><tbody>"
            "<tr><th>Field</th><th>Rule</th></tr>"
            "<tr><td>field_a</td><td>rule_a</td></tr>"
            "</tbody></table>"
        )
        lines = [l for l in self.convert(html).split("\n") if l.strip()]
        self.assertTrue(all(l.startswith("|") for l in lines), lines)
        self.assertIn("| field_a | rule_a |", lines)

    def test_image_alt_and_filename_survive(self):
        html = ('<ac:image ac:alt="Retry diagram">'
                '<ri:attachment ri:filename="retry-flow.png"/></ac:image>')
        out = self.convert(html)
        self.assertIn("Retry diagram", out)
        self.assertIn("retry-flow.png", out)

    def test_anchor_text_survives_without_the_url(self):
        out = self.convert('<p>See <a href="https://example.invalid/spec">the spec</a>.</p>')
        self.assertIn("the spec", out)
        self.assertNotIn("example.invalid", out)


@unittest.skipUnless(_has("docx"), "python-docx not installed")
class DocxRealLibraryTests(unittest.TestCase):
    """`doc.paragraphs` excludes table paragraphs — the bug a stub cannot show."""

    def _build(self):
        from docx import Document

        doc = Document()
        doc.add_heading("Requirements", level=1)
        doc.add_paragraph("Lead-in prose.")
        table = doc.add_table(rows=3, cols=3)
        values = [
            ["Requirement", "Priority", "Owner"],
            ["req_alpha", "P1", "team_a"],
            ["req_beta", "P2", "team_b"],
        ]
        for row, vals in zip(table.rows, values):
            for cell, val in zip(row.cells, vals):
                cell.text = val
        doc.add_paragraph("Trailing prose.")

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def test_iter_inner_content_exists_on_the_installed_version(self):
        from docx import Document
        self.assertTrue(
            hasattr(Document(), "iter_inner_content"),
            "document-order walking requires python-docx >= 1.1",
        )

    def test_table_content_is_extracted(self):
        from ingestion.file_ingestor import convert_file_to_markdown
        out = convert_file_to_markdown("spec.docx", self._build())
        self.assertIn("| req_alpha | P1 | team_a |", out)
        self.assertIn("| req_beta | P2 | team_b |", out)

    def test_table_has_a_header_and_delimiter(self):
        from ingestion.file_ingestor import convert_file_to_markdown
        out = convert_file_to_markdown("spec.docx", self._build())
        self.assertIn("| Requirement | Priority | Owner |", out)
        self.assertIn("| --- | --- | --- |", out)

    def test_prose_and_table_keep_document_order(self):
        from ingestion.file_ingestor import convert_file_to_markdown
        out = convert_file_to_markdown("spec.docx", self._build())
        self.assertLess(out.index("Lead-in prose."), out.index("| req_alpha"))
        self.assertLess(out.index("| req_beta"), out.index("Trailing prose."))

    def test_heading_becomes_a_markdown_heading(self):
        from ingestion.file_ingestor import convert_file_to_markdown
        out = convert_file_to_markdown("spec.docx", self._build())
        self.assertIn("# Requirements", out)


@unittest.skipUnless(_has("openpyxl"), "openpyxl not installed")
class ExcelRealLibraryTests(unittest.TestCase):
    def _build(self, merge=False, leading_blank=False):
        import openpyxl

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Rules"
        rows = [
            ["Module", "Requirement", "Priority"],
            ["module_x", "req_alpha", "P1"],
            [None, "req_beta", "P2"],
        ]
        if leading_blank:
            rows.insert(0, [None, None, None])
        for row in rows:
            ws.append(row)
        if merge:
            base = 3 if leading_blank else 2
            ws.merge_cells(start_row=base, start_column=1, end_row=base + 1, end_column=1)
        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()

    def test_sheet_becomes_a_markdown_table(self):
        from ingestion.file_ingestor import convert_file_to_markdown
        out = convert_file_to_markdown("rules.xlsx", self._build())
        self.assertIn("| Module | Requirement | Priority |", out)
        self.assertIn("| --- | --- | --- |", out)

    def test_leading_blank_row_still_yields_a_header(self):
        """The old `i == 0` header check emitted no delimiter at all here."""
        from ingestion.file_ingestor import convert_file_to_markdown
        out = convert_file_to_markdown("rules.xlsx", self._build(leading_blank=True))
        self.assertIn("| Module | Requirement | Priority |", out)
        self.assertIn("| --- | --- | --- |", out)

    def test_merged_cell_is_filled_down(self):
        from ingestion.file_ingestor import convert_file_to_markdown
        out = convert_file_to_markdown("rules.xlsx", self._build(merge=True))
        self.assertIn("| module_x | req_beta | P2 |", out,
                      "a merged Module cell must repeat on the covered row")

    def test_sheet_name_becomes_a_section(self):
        from ingestion.file_ingestor import convert_file_to_markdown
        self.assertIn("## Rules", convert_file_to_markdown("rules.xlsx", self._build()))


@unittest.skipUnless(_has("tiktoken") and _has("numpy"), "chunker deps not installed")
class ChunkerRealTokenizerTests(unittest.TestCase):
    """Token accounting with the real cl100k_base encoding, not a chars//4 stub."""

    def _table(self, n):
        head = "| Field | Rule | Owner |\n| --- | --- | --- |"
        rows = "\n".join(f"| field_{i} | rule_{i} | team_{i % 3} |" for i in range(n))
        return f"{head}\n{rows}"

    def test_long_table_splits_and_every_chunk_keeps_the_header(self):
        from ingestion.chunker import chunk_document

        chunks = chunk_document(f"# D\n\n## Rules\n\n{self._table(400)}\n", "t:int")
        self.assertGreater(len(chunks), 1)
        for c in chunks:
            self.assertTrue(c["chunk_text"].startswith("| Field | Rule | Owner |"))

    def test_yaml_separators_survive_real_tokenization(self):
        from ingestion.chunker import chunk_document

        doc = "# D\n\n## Config\n\n```yaml\n---\nkey: value\n---\nnext: thing\n```\n"
        self.assertEqual(chunk_document(doc, "t:int3")[0]["chunk_text"].count("---"), 2)


@unittest.skipUnless(
    _has("tiktoken") and _has("numpy") and _has("sentence_transformers"),
    "chunker or embedding deps not installed",
)
class EmbeddingCapAlignmentTests(unittest.TestCase):
    """
    The embedder's character cap must stay above the largest chunk the chunker
    emits. When they drift apart the tail of a chunk is dropped from the vector
    while Elasticsearch still indexes it whole for BM25 — dense and keyword
    retrieval then see different documents, with no error anywhere.

    Gated on sentence_transformers too: importing embed_client loads it.
    """

    def test_cap_exceeds_the_chunkers_maximum(self):
        from ingestion.chunker import SEMANTIC_MAX
        from embeddings.embed_client import EMBED_MAX_CHARS

        # ~4 chars per token for cl100k_base on prose, plus the heading prefix
        # and a carried table header.
        self.assertGreater(EMBED_MAX_CHARS, SEMANTIC_MAX * 4,
                           "QA_EMBED_MAX_CHARS is below the chunker's maximum chunk")

    def test_real_chunks_fit_under_the_cap(self):
        from ingestion.chunker import chunk_document
        from embeddings.embed_client import EMBED_MAX_CHARS

        head = "| Field | Rule | Owner |\n| --- | --- | --- |"
        rows = "\n".join(f"| field_{i} | rule_{i} | team_{i % 3} |" for i in range(2000))
        chunks = chunk_document(f"# D\n\n## Rules\n\n{head}\n{rows}\n", "t:cap")
        for c in chunks:
            formatted = f"Section: {c['section_heading']}\n\n{c['chunk_text']}"
            self.assertLessEqual(len(formatted), EMBED_MAX_CHARS)


if __name__ == "__main__":
    unittest.main()
